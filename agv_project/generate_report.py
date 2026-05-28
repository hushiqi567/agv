"""生成开题报告差距分析 .docx 文件"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)

# ===== 标题 =====
title = doc.add_heading('开题报告 vs 当前项目实现 — 差距分析报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('项目：基于强化学习的AGV路径规划算法设计')
doc.add_paragraph('分析日期：2026年5月26日')
doc.add_paragraph()

# ===== 一 =====
doc.add_heading('一、当前项目实际完成情况', level=1)
doc.add_paragraph('当前项目实现了一个 50×50 仓库网格中的多AGV仿真系统，包含以下模块：')

for item in [
    'env/：仓库环境（50×50网格，20个进货口，12个出货口，10个随机移动障碍物），Pygame可视化渲染',
    'scheduler/：任务调度（泊松分布生成任务，最近距离优先分配，OD流程管理）',
    'path_planning/：路径规划（CBS算法全局规划 + DQN算法局部避撞 + AGV控制器）',
    'interface/：接口定义（发布-订阅消息总线、全局配置管理、数据类型定义）',
    'main.py：仿真主入口，支持命令行参数，含RL训练模式',
    'rl_model.pth / rl_model_v2.pth：预训练的RL模型权重文件（各约1.7MB）',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# ===== 二 =====
doc.add_heading('二、最核心的偏差：RL 的角色定位', level=1)
doc.add_paragraph(
    '开题报告的标题和全文核心主张是"基于强化学习的AGV路径规划"，'
    '即 RL 应当是路径规划的主角。但当前项目的实际架构是：'
)
doc.add_paragraph('CBS 算法（传统冲突搜索）负责全局路径规划', style='List Bullet')
doc.add_paragraph('DQN 算法（深度Q网络）仅负责 5×5 局部网格内的避撞', style='List Bullet')
doc.add_paragraph('结论：RL 只是辅助角色，不是路径规划的核心', style='List Bullet')

table1 = doc.add_table(rows=4, cols=3)
table1.style = 'Light Grid Accent 1'
for i, text in enumerate(['开题报告承诺', '实际实现', '偏差程度']):
    table1.rows[0].cells[i].text = text
for i, row_data in enumerate([
    ['"基于强化学习的AGV路径规划算法设计"', 'CBS(传统搜索)做全局规划 + RL做局部避撞', '严重——RL不是主角'],
    ['"AGV可自主感知环境变化、动态调整路径"', '路径调整靠CBS每步完全重算，非RL动态决策', '严重——核心机制不同'],
    ['"AGV能通过试错-反馈实时交互，在线学习最优路径"', '在线学习只发生在局部避撞层面，全局路径靠A*+CBS', '严重——学习范围极度受限'],
]):
    for j, text in enumerate(row_data):
        table1.rows[i+1].cells[j].text = text

doc.add_paragraph()

# ===== 三 =====
doc.add_heading('三、完全缺失的功能（零实现）', level=1)

missing = [
    ('多智能体强化学习 (MARL) 框架',
     '报告研究内容第4条要求"基于多智能体强化学习（MARL）框架，设计多AGV协同机制"。'
     '当前每个AGV各跑各的DQN，无任何智能体间通信或协同决策。'),
    ('死锁检测与预防机制',
     '报告立项意义中明确提到"避免多个AGV发生拥堵甚至死锁现象"。'
     '当前无任何死锁检测、预防或恢复逻辑。'),
    ('多目标优化的奖励函数',
     '报告承诺将"能耗成本、装卸货等待时间"融入奖励函数。'
     '当前奖励函数仅6项（到达目标+100、完成任务+150、碰撞-50、距离变化±1、步数惩罚-0.1、等待惩罚-0.5），'
     '缺失能耗、装卸等待时间、优先级权重、通道拥堵度。'),
    ('任务优先级差异化调度',
     '报告创新点提到"结合订单优先级等因素，设计差异化奖励函数"。'
     'Task.priority字段始终硬编码为1，调度仅用"最近距离优先"。'),
    ('技术报告',
     '预期提交成果第2项。当前仅有约150行的README.md，无技术报告。'),
    ('对比实验与消融实验',
     '报告研究路线阶段4要求"多轮测试，记录指标，根据结果优化"。'
     '当前无任何对比实验：RL vs 传统算法、RL vs 无RL、单智能体 vs 多智能体等。'),
    ('文献综述',
     '报告研究路线阶段1要求"调研AGV路径规划相关文献，形成文献综述"。'
     '当前无任何文献综述文件。'),
]

for title_text, detail in missing:
    p = doc.add_paragraph()
    r = p.add_run(f'[缺失] {title_text}')
    r.bold = True
    doc.add_paragraph(f'    要求：{detail}')

doc.add_paragraph()

# ===== 四 =====
doc.add_heading('四、部分实现但不达标的', level=1)

table2 = doc.add_table(rows=8, cols=4)
table2.style = 'Light Grid Accent 1'
for i, text in enumerate(['开题报告要求', '当前实现', '不足', '严重程度']):
    table2.rows[0].cells[i].text = text
for i, row_data in enumerate([
    ['状态空间：位置、障碍物', '5×5局部网格(3通道)', '50×50地图中视野极小；不包含装载状态、电量、拥堵度', '严重'],
    ['动作空间：前进、转向', '上/下/左/右/等待（网格移动）', '无车辆运动学模型，不符合"前进、转向"描述', '中等'],
    ['仿真环境：动态障碍物+多AGV', '50×50仓库+10随机移动障碍物+8AGV', '无充电站、无可配置布局、无拥堵场景生成', '中等'],
    ['性能指标：路径长度、运输耗时、避障成功率', '步数、完成任务数、AGV状态计数', '完全缺失承诺的三项核心指标', '严重'],
    ['多轮测试+参数优化', 'test函数仅为基本smoke test', '无多轮测试、无参数调优记录', '严重'],
    ['先单智能体再扩展多智能体', '直接做了多AGV（且用CBS非RL）', '跳过单智能体RL基线，无法论证多智能体协同必要性', '严重'],
    ['算法与环境对接', '模块间方法直接调用', '消息总线架构被架空，未真正使用', '中等'],
]):
    for j, text in enumerate(row_data):
        table2.rows[i+1].cells[j].text = text

doc.add_paragraph()

# ===== 五 =====
doc.add_heading('五、其他代码与工程层面的问题', level=1)

for issue in [
    '命名体系混乱：进货口/出货口的英文变量名(LOADING/UNLOADING)与中文注释互相矛盾，多处标注不一致',
    '数据结构重复定义：MapConfig、SimulationConfig、manhattan_distance 在两个以上文件中重复定义',
    'AGV状态多处独立维护：TaskAllocator 和 AGVController 各自维护AGV状态字典，互不同步',
    '消息总线被架空：定义了完整的发布-订阅模式，但实际运行中全部绕开，模块间直接调用',
    'CBS最大迭代次数不足(100次)：8台AGV在50×50地图上通常需要远多于100次迭代才能找到无冲突解',
    'RL经验收集有时序bug：_collect_path_move_experience()中获取"移动前"状态时AGV位置已被更新',
    'get_state() 在 use_numpy=False 时行为不一致（第398行逻辑bug）',
    '硬编码散落各处：NUM_AGVS=8、障碍物数量=10、AGV初始位置直接写在代码中，config中对应配置未被使用',
    '无 .gitignore：__pycache__/ 和 .pyc 文件被提交到git仓库',
    '大二进制文件在git中：rl_model.pth(1.7MB) 和 rl_model_v2.pth(1.7MB)',
    '无单元测试：整个项目零测试代码',
    'sys.path 操作出现在每个文件中',
    'pygame 无条件导入：即使 --no-render 模式也需要安装 pygame',
    '日志重复配置：ConfigManager 和 main() 分别配置了日志handler',
    '充电站仅定义枚举值(CellType.CHARGING_STATION=4)，从未使用',
    'AGVConfig 中电量相关参数齐全但从未被任何逻辑使用',
]:
    doc.add_paragraph(issue, style='List Bullet')

doc.add_paragraph()

# ===== 六 =====
doc.add_heading('六、开题报告研究内容逐条对照', level=1)

table3 = doc.add_table(rows=6, cols=4)
table3.style = 'Light Grid Accent 1'
for i, text in enumerate(['报告研究内容', '对应实现', '完成度', '说明']):
    table3.rows[0].cells[i].text = text
for i, row_data in enumerate([
    ['1. 学习RL理论、线性代数、Python编程等基础工具', 'N/A', '未体现', '代码使用了PyTorch和NumPy，但无学习笔记或文献综述'],
    ['2. 调研AGV路径规划行业现状，梳理传统方法优缺点', 'N/A', '未体现', '无文献综述文件；代码用了CBS(传统方法)但未做系统性梳理'],
    ['3. 设计状态/动作/奖励函数，实现单AGV最优路径规划', 'rl_collision_avoidance.py', '约30%', '状态过于简化(5×5局部网格)，奖励函数未含能耗/等待时间，RL只做局部避撞而非完整路径规划'],
    ['4. 基于MARL框架，设计多AGV协同机制', '不存在', '0%', '无任何MARL实现，多AGV协调完全依赖CBS算法'],
    ['5. 仿真环境：构建仓储场景，配置AGV参数、动态元素', 'env/ + renderer.py', '约60%', '基础环境存在，但无充电站、无可配置布局、无拥堵模拟、无指标采集'],
]):
    for j, text in enumerate(row_data):
        table3.rows[i+1].cells[j].text = text

doc.add_paragraph()

# ===== 七 =====
doc.add_heading('七、预期提交成果达成情况', level=1)

table4 = doc.add_table(rows=3, cols=3)
table4.style = 'Light Grid Accent 1'
for i, text in enumerate(['预期成果', '完成度', '差距']):
    table4.rows[0].cells[i].text = text
for i, row_data in enumerate([
    ['成果1：搭建适配智能无人仓储场景的RL仿真环境', '约50%',
     '基础环境已搭建(50×50网格+动态障碍物+8AGV)，但无MARL协同、无充电站、无拥堵模拟、无指标采集、无可配置场景'],
    ['成果2：形成1份完整技术报告', '约5%',
     '仅有README.md。缺失：研究背景、传统方法梳理、RL算法原理详解、测试结果(缺指标数据)、效率对比分析'],
]):
    for j, text in enumerate(row_data):
        table4.rows[i+1].cells[j].text = text

doc.add_paragraph()

# ===== 八 =====
doc.add_heading('八、优先修复建议（按紧急程度排序）', level=1)

priorities = [
    ('P0 — 明确RL的角色定位',
     '要么将RL升级为真正的路径规划核心（取代CBS做全局规划），要么修改项目描述为"CBS+RL混合路径规划"。'
     '当前状态与开题报告标题直接矛盾，是答辩时最致命的硬伤。'),
    ('P0 — 实现MARL协同机制',
     '这是研究内容第4条的核心承诺，也是创新点的主要支撑。至少需要一个基础的MARL框架（如MADDPG、QMIX等）。'),
    ('P1 — 补全核心指标采集系统',
     '必须能采集：每条路径的长度、每个任务的完成耗时、避障成功率/碰撞次数。这是撰写技术报告和做对比实验的前提。'),
    ('P1 — 构建单智能体RL基线',
     '按研究路线要求，先完成单AGV RL路径规划，作为与多AGV方案对比的基线。这是论证MARL必要性的前提。'),
    ('P2 — 扩展奖励函数',
     '加入能耗成本、装卸等待时间、任务优先级等因素，体现多目标优化能力。'),
    ('P2 — 添加死锁检测',
     '报告中明确提到的问题，需要有基本的检测和处理机制。'),
    ('P3 — 修复代码工程问题',
     '统一命名、添加.gitignore、移除二进制文件、收敛配置、补充测试等。'),
    ('P3 — 撰写技术报告',
     '两份预期成果之一，必须完成。建议在补完指标采集和对比实验后再写。'),
]

for title_text, detail in priorities:
    p = doc.add_paragraph()
    r = p.add_run(title_text)
    r.bold = True
    doc.add_paragraph(f'    {detail}')

doc.add_paragraph()
doc.add_paragraph('— 报告完 —')

doc.save('开题报告差距分析.docx')
print('Done: 开题报告差距分析.docx')
